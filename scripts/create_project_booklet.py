from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\w10_dashboard")
OUTPUT_DIR = ROOT / "docs" / "deliverables"
OUTPUT_PATH = OUTPUT_DIR / "W10-Dashboard-technical-booklet.docx"
LOGO_PATH = ROOT / "public" / "picture" / "egat.png"
W10_ICON_PATH = ROOT / "public" / "picture" / "รูปภาพ14-Photoroom.png"


COLORS = {
    "navy": "0F2747",
    "blue": "005B9A",
    "blue_dark": "003C66",
    "amber": "F0B323",
    "green": "1F7A4D",
    "rose": "B42318",
    "orange": "B54708",
    "mist": "F2F6FA",
    "muted": "E8EEF4",
    "border": "D8E1EA",
    "ink": "17202A",
    "gray": "52606D",
    "light_blue": "EAF3F9",
    "light_amber": "FFF8E5",
    "light_green": "EAF6EF",
    "light_rose": "FFF0EF",
    "white": "FFFFFF",
}


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(run, name: str = "Calibri", size: float | None = None,
                 color: str | None = None, bold: bool | None = None,
                 italic: bool | None = None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    # Thai glyphs use the Windows fallback while Latin text follows the preset.
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name: str = "Calibri", size: float | None = None,
                   color: str | None = None, bold: bool | None = None):
    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial")
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = rgb(color)
    if bold is not None:
        style.font.bold = bold


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top: int = 80, start: int = 120,
                     bottom: int = 80, end: int = 120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")
    cell.width = Inches(width_dxa / 1440)


def set_table_geometry(table, widths_dxa: Sequence[int], indent_dxa: int = 120,
                       borders: str = "single"):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    borders_node = tbl_pr.find(qn("w:tblBorders"))
    if borders_node is None:
        borders_node = OxmlElement("w:tblBorders")
        tbl_pr.append(borders_node)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders_node.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders_node.append(node)
        node.set(qn("w:val"), borders if borders != "none" else "nil")
        if borders == "single":
            node.set(qn("w:sz"), "4")
            node.set(qn("w:space"), "0")
            node.set(qn("w:color"), COLORS["border"])

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_keep_with_next(paragraph, value: bool = True):
    p_pr = paragraph._p.get_or_add_pPr()
    node = p_pr.find(qn("w:keepNext"))
    if node is None:
        node = OxmlElement("w:keepNext")
        p_pr.append(node)
    node.set(qn("w:val"), "1" if value else "0")


def add_page_field(run):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    separate.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instruction)
    run._r.append(separate)
    run._r.append(end)


def set_paragraph_border(paragraph, color: str = "D8E1EA", size: str = "8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None,
             color: str | None = None, italic: bool = False,
             after: float = 8, before: float = 0,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.333
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, size=11, color=color or COLORS["ink"], bold=True, italic=italic)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest, size=11, color=color or COLORS["ink"], italic=italic)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=11, color=color or COLORS["ink"], italic=italic)
    return paragraph


def add_bullet(doc: Document, text: str, level: int = 0):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.194)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.208
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=COLORS["ink"])
    return paragraph


def add_number(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.194)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.208
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=COLORS["ink"])
    return paragraph


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, size={1: 16, 2: 13, 3: 12}[level],
                 color=COLORS["blue"] if level < 3 else COLORS["blue_dark"],
                 bold=True)
    return paragraph


def add_kicker(doc: Document, text: str, color: str = "B54708"):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run(text.upper())
    set_run_font(run, size=10, color=color, bold=True)
    return paragraph


def add_callout(doc: Document, label: str, text: str,
                fill: str = "F4F6F9", accent: str = "005B9A"):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent_dxa=120, borders="none")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    label_run = p.add_run(label)
    set_run_font(label_run, size=10.5, color=accent, bold=True)
    body_run = p.add_run("  " + text)
    set_run_font(body_run, size=10.5, color=COLORS["ink"])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_table(doc: Document, headers: Sequence[str], rows: Iterable[Sequence[str]],
              widths_dxa: Sequence[int], header_fill: str = "E8EEF4",
              font_size: float = 9.5, first_col_bold: bool = False):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa, indent_dxa=120, borders="single")
    header_row = table.rows[0]
    mark_header_row(header_row)
    for index, text in enumerate(headers):
        cell = header_row.cells[index]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        set_run_font(run, size=font_size, color=COLORS["navy"], bold=True)

    for row_index, row_data in enumerate(rows):
        cells = table.add_row().cells
        for index, text in enumerate(row_data):
            cell = cells[index]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(text))
            set_run_font(run, size=font_size, color=COLORS["ink"], bold=first_col_bold and index == 0)
            if row_index % 2 == 1:
                set_cell_shading(cell, "FAFCFD")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_code_block(doc: Document, lines: Sequence[str]):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent_dxa=120, borders="none")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F9FB")
    cell.text = ""
    for index, line in enumerate(lines):
        p = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        run = p.add_run(line)
        set_run_font(run, name="Consolas", size=9, color=COLORS["navy"])
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Consolas")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_source_note(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run("แหล่งข้อมูล: " + text)
    set_run_font(run, size=8.5, color=COLORS["gray"], italic=True)
    return p


def add_section_break(doc: Document):
    paragraph = doc.add_page_break()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.keep_with_next = False
    return paragraph


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_style_font(normal, size=11, color=COLORS["ink"])
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.333

    for level, size, color, before, after in (
        (1, 16, COLORS["blue"], 18, 10),
        (2, 13, COLORS["blue"], 12, 6),
        (3, 12, COLORS["blue_dark"], 8, 4),
    ):
        style = doc.styles[f"Heading {level}"]
        set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        set_style_font(style, size=11, color=COLORS["ink"])
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    # Header/footer are intentionally quiet so the cover remains editorial.
    header = section.header
    header_p = header.paragraphs[0]
    header_p.text = ""
    header_p.paragraph_format.space_after = Pt(2)
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left = header_p.add_run("W10 Dashboard  |  Technical Project Booklet")
    set_run_font(left, size=8.5, color=COLORS["gray"], bold=True)
    set_paragraph_border(header_p, COLORS["border"], "4")

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.text = ""
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(3)
    footer_run = footer_p.add_run("W10 Dashboard  |  Internal documentation  |  หน้า ")
    set_run_font(footer_run, size=8.5, color=COLORS["gray"])
    page_run = footer_p.add_run()
    set_run_font(page_run, size=8.5, color=COLORS["gray"])
    add_page_field(page_run)


def add_cover(doc: Document):
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)

    if LOGO_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        inline = run.add_picture(str(LOGO_PATH), width=Inches(1.05))
        inline._inline.docPr.set('descr', 'ตราสัญลักษณ์ EGAT')
        inline._inline.docPr.set('title', 'EGAT logo')

    add_kicker(doc, "PROJECT BOOKLET  /  TECHNICAL & OPERATIONS REVIEW")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("W10 Dashboard")
    set_run_font(run, size=30, color=COLORS["navy"], bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("รูปเล่มวิเคราะห์ระบบและคู่มือปฏิบัติการ")
    set_run_font(run, size=15, color=COLORS["blue"], bold=True)

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule.paragraph_format.space_after = Pt(20)
    set_paragraph_border(rule, COLORS["amber"], "12")

    metadata = [
        ("ขอบเขต", "โค้ด แอปพลิเคชัน API data layer tests scripts และเอกสารโครงการ"),
        ("สถานะอ้างอิง", "repository snapshot วันที่ 31 กรกฎาคม 2026"),
        ("commit ล่าสุด", "0955f46  |  main  |  fix(ui): polish ot employee console chrome"),
        ("ผู้อ่านหลัก", "ผู้บริหารหน่วยงาน ผู้ดูแลระบบ นักพัฒนา และผู้ปฏิบัติงาน W10"),
    ]
    add_table(doc, ["รายการ", "รายละเอียด"], metadata, [1800, 7560],
              header_fill=COLORS["light_blue"], font_size=9.5, first_col_bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(42)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("เอกสารนี้สรุปจากระบบจริงโดยไม่เปิดเผยค่า secret, token หรือ private key")
    set_run_font(run, size=9, color=COLORS["gray"], italic=True)
    doc.add_page_break()


def add_contents(doc: Document):
    add_kicker(doc, "READING MAP", COLORS["blue"])
    add_heading(doc, "สารบัญและวิธีอ่าน", 1)
    add_body(doc, "รูปเล่มนี้แบ่งเป็น 11 ส่วน เริ่มจากภาพรวมของระบบ ไปสู่การอ่านแต่ละ route, สถาปัตยกรรมข้อมูล, Shop Order แบบเจาะลึก, ผลตรวจคุณภาพ และข้อเสนอแนะสำหรับการดูแลระบบใน Production.")
    sections = [
        ("1", "บทสรุปผู้บริหาร", "ระบบทำอะไร เหมาะกับใคร และผลตรวจปัจจุบัน"),
        ("2", "ขอบเขตและเทคโนโลยี", "โครงสร้าง repository, stack และขอบเขตการอ่าน"),
        ("3", "เส้นทางใช้งานและ UI shell", "7 routes, navigation และสถานะการ migrate"),
        ("4", "การทำงานของแต่ละหน้า", "Home, Purchasing, BEML, OT และ Shop Order"),
        ("5", "สถาปัตยกรรมและแหล่งข้อมูล", "Browser, Next.js API, Google Sheets/Drive"),
        ("6", "Shop Order เชิงลึก", "CRUD, schema A-K, upload และ lifecycle"),
        ("7", "ความปลอดภัยและความทนทาน", "validation, OAuth, WAF, logging และ failure modes"),
        ("8", "การทดสอบและผลตรวจ", "tests, TypeScript, build, lint และ browser handoff"),
        ("9", "ข้อค้นพบและงานค้าง", "ความเสี่ยงและ technical debt ที่พบจากการอ่าน"),
        ("10", "คู่มือ deploy และ runbook", "ตั้งค่าเครื่อง, Vercel, OAuth และ cleanup"),
        ("11", "ภาคผนวก", "API reference, file map, glossary และแหล่งข้อมูล"),
    ]
    add_table(doc, ["ส่วน", "หัวข้อ", "คำตอบที่ผู้อ่านจะได้"], sections,
              [700, 2700, 5960], header_fill=COLORS["light_blue"], font_size=9.4)
    add_callout(doc, "คำแนะนำ", "ถ้าต้องการทำความเข้าใจเร็ว ให้เริ่มที่ส่วน 1, 3, 5 และ 6; ถ้ารับผิดชอบ Production ให้ต่อด้วยส่วน 7, 8, 9 และ 10.", fill=COLORS["light_amber"], accent=COLORS["orange"])


def add_executive_summary(doc: Document):
    add_section_break(doc)
    add_kicker(doc, "01  /  EXECUTIVE SUMMARY")
    add_heading(doc, "บทสรุปผู้บริหาร", 1)
    add_body(doc, "W10 Dashboard คือแดชบอร์ดปฏิบัติการภายในสำหรับทีมซ่อมบำรุงและงานธุรการของ W10 ใช้ Google Sheets เป็นแหล่งข้อมูลหลัก แล้วนำเสนอเป็นหน้าจอที่อ่านเร็วสำหรับติดตาม Work Order, สถานะจัดซื้อจัดจ้าง, คลังอะไหล่ BEML และ OT โดยแยกมุมมองพนักงานกับลูกจ้างอย่างชัดเจน.")
    add_body(doc, "จุดเด่นของระบบอยู่ที่การรวมข้อมูลจากหลายชีตไว้ใน route ที่ใช้งานซ้ำทุกวัน และการออกแบบ Shop Order รุ่น Next.js ให้มีการตรวจไฟล์แบบหลายชั้น, resumable upload, OAuth Drive แบบ scope ต่ำ, partial success เมื่อไฟล์แนบล้มเหลว และ cleanup ที่ย้ายไฟล์เข้า Trash อย่างหน่วงเวลา.")
    add_callout(doc, "ข้อสรุปสั้น", "แกนธุรกิจและการทดสอบของ Shop Order มีความเป็นระบบสูง แต่ทั้ง repository ยังมี lint baseline ที่ไม่ผ่าน, การย้าย UI shell ยังอยู่เพียง pilot route และยังมี public-access boundary ที่ต้องบริหารด้วย URL/WAF.", fill=COLORS["light_blue"], accent=COLORS["blue"])

    add_heading(doc, "ตัวเลขสำคัญจาก snapshot นี้", 2)
    facts = [
        ("7", "หน้าใช้งานหลัก", "/, /purchasing, /purchasing-all, /beml-inventory, /ot-summary, /ot-employee, /shop-order"),
        ("12", "API route handler", "รวม route หลักของ dashboard, OT, purchasing, inventory และ Shop Order"),
        ("25", "test files", "ผลรันจริงของ Vitest"),
        ("270", "tests passed", "unit, component, API และ repository-style tests"),
        ("10 MB", "ขนาดไฟล์แนบสูงสุด", "JPEG, PNG, WebP และ PDF สำหรับ Shop Order"),
        ("30 วัน", "ระยะหน่วงก่อนย้ายไฟล์เก่าเข้า Trash", "เฉพาะไฟล์ที่ OAuth ของระบบเป็นเจ้าของ"),
    ]
    add_table(doc, ["ค่า", "ความหมาย", "รายละเอียด"], facts,
              [1000, 2400, 5960], header_fill=COLORS["light_green"], font_size=9.4, first_col_bold=True)

    add_heading(doc, "ผลตรวจที่ยืนยันได้", 2)
    checks = [
        ("PASS", "npm run test:unit", "25 files / 270 tests"),
        ("PASS", "npm run test:oauth-setup", "5 tests"),
        ("PASS", "npx tsc --noEmit", "ไม่มี output error"),
        ("PASS", "npm run build", "Next.js 16.2.12, production build สำเร็จ"),
        ("PASS", "scoped ESLint", "shell/navigation/route pilot ผ่าน"),
        ("FAIL", "npm run lint", "144 errors / 10 warnings ทั้ง repository"),
        ("PENDING", "browser manual acceptance", "มี HTTP smoke/build baseline แต่ viewport/focus/console checks ยังรอ"),
    ]
    add_table(doc, ["สถานะ", "คำสั่งหรือกิจกรรม", "ผล"], checks,
              [1000, 3300, 5060], header_fill=COLORS["light_amber"], font_size=8.6)


def add_scope_stack(doc: Document):
    add_kicker(doc, "02  /  SCOPE & STACK", COLORS["blue"])
    add_heading(doc, "ขอบเขตการอ่านและเทคโนโลยี", 1)
    add_body(doc, "การวิเคราะห์ครอบคลุม source code ใน app, components, lib และ scripts; test files; เอกสาร README, PRODUCT, CONTEXT, specs, plans และ verification handoff; ตลอดจน configuration ที่เกี่ยวข้อง. ไม่อ่านเนื้อหาใน node_modules และ .next เพราะเป็น dependency/build output และไม่เปิดเผยค่าใน .env.local.")
    scope_rows = [
        ("Runtime", "Next.js 16.2.12 App Router, React 19.2.4, TypeScript strict"),
        ("Styling", "Tailwind CSS 4, app/globals.css, Prompt ผ่าน next/font/google"),
        ("Interaction", "Framer Motion, Lucide React, custom responsive dialogs"),
        ("Charts", "Highcharts + highcharts-more, Recharts, react-d3-speedometer"),
        ("Data", "googleapis, Google Sheets Service Account, Google Drive OAuth drive.file"),
        ("Testing", "Vitest 4.1.10, Testing Library, jsdom, Node test runner"),
        ("Deploy", "Vercel, vercel.json cron สำหรับ /api/shop-order/cleanup"),
    ]
    add_table(doc, ["ชั้นระบบ", "เทคโนโลยีและบทบาท"], scope_rows,
              [1800, 7560], header_fill=COLORS["light_blue"], font_size=9.5, first_col_bold=True)

    add_heading(doc, "โครงสร้าง repository แบบอ่านเร็ว", 2)
    tree_rows = [
        ("app/", "routes, page entry, API route handlers, global layout/style"),
        ("components/", "navigation, AppShell, page header, charts และ Shop Order UI"),
        ("lib/", "Google Sheets adapter และ domain/repository ของ Shop Order"),
        ("scripts/", "OAuth setup utility และ Node test"),
        ("docs/superpowers/", "design specs, implementation plans, verification handoff"),
        ("public/picture/", "โลโก้/ภาพประกอบของ W10 และทีมงาน"),
        ("scratch/", "เครื่องมือ debug/ตรวจชีตเก่า ไม่ใช่ runtime production"),
    ]
    add_table(doc, ["โฟลเดอร์", "หน้าที่"], tree_rows,
              [2200, 7160], header_fill=COLORS["muted"], font_size=9.4, first_col_bold=True)
    add_callout(doc, "หลักการอ่าน", "ค่าต่าง ๆ ที่มาจาก Google Sheets ถูก map ด้วยตำแหน่งคอลัมน์/แถวที่กำหนดใน code; ดังนั้นชื่อ tab, range และโครงสร้างชีตเป็น dependency เชิงปฏิบัติการ ไม่ใช่เพียงรายละเอียดการตั้งค่า.", fill=COLORS["light_amber"], accent=COLORS["orange"])

    add_heading(doc, "Product intent", 2)
    add_bullet(doc, "ผู้ใช้หลักคือเจ้าหน้าที่ W10 maintenance และ administrative staff ที่ต้องตรวจยอดและสถานะงานระหว่างวัน.")
    add_bullet(doc, "คำศัพท์ Employee, Contractor, OT Summary และ ETAS Scan Data ถูกกำหนดไว้ใน CONTEXT.md เพื่อไม่ให้มุมมองข้อมูลคนละประเภทปะปนกัน.")
    add_bullet(doc, "แนวทาง UI คือ operational, direct, familiar; หลีกเลี่ยง marketing-style landing page และ motion ที่ทำให้การอ่านช้าลง.")


def add_routes_shell(doc: Document):
    add_kicker(doc, "03  /  ROUTES & SHELL", COLORS["blue"])
    add_heading(doc, "เส้นทางใช้งานและ UI shell", 1)
    add_body(doc, "ระบบมี canonical navigation ที่รวมไว้ใน components/navigation/navigationDestinations.ts และใช้ซ้ำทั้ง dropdown แบบเดิมกับ AppShell ใหม่. NavigationMenu รองรับ hover, click/touch, outside pointer, Escape, focus return, current-route state และ reduced motion.")
    route_rows = [
        ("/", "W10 Dashboard", "ยอด W/O, status, group, equipment, load factor", "legacy"),
        ("/purchasing", "การจัดซื้อจัดจ้าง", "สรุปหมวด, status, gauges, purchase table", "legacy"),
        ("/purchasing-all", "สถานะการซื้อจ้างทั้งหมด", "ภาพรวมทั้งหมด ใช้ purchasing content แบบ fixed filters", "legacy"),
        ("/beml-inventory", "คลังอะไหล่ BEML", "stock map, health ring, inventory table/modal", "legacy"),
        ("/ot-summary", "สรุป OT ลูกจ้าง", "OT ลูกจ้าง, ETAS scan, error tables", "legacy"),
        ("/ot-employee", "สรุป OT พนักงาน", "OT พนักงาน, ETAS scan, error tables", "console pilot"),
        ("/shop-order", "Shop Order", "CRUD, filter, summary, attachment workflow", "legacy"),
    ]
    add_table(doc, ["Route", "ชื่อหน้า", "หน้าที่", "Chrome ปัจจุบัน"], route_rows,
              [1500, 2100, 4400, 1360], header_fill=COLORS["light_blue"], font_size=8.8)

    add_heading(doc, "ภาพรวมการ migrate shell", 2)
    add_body(doc, "ShellMigrationGate ใน app/layout.tsx อ่าน pathname แล้วห่อเฉพาะ route ที่อยู่ใน consoleRoutes. ณ snapshot นี้ consoleRoutes มีเพียง /ot-employee; route อื่นยัง render children แบบเดิม. นี่เป็นการ rollout แบบ phased ตามแผน ไม่ใช่การย้ายครบทั้งระบบ.")
    shell_rows = [
        ("Shared shell", "Sidebar desktop 240px, MobileTopBar, MobileNavigationDrawer"),
        ("PageHeader", "title, description, sync status, filters/actions slot, refresh state"),
        ("RouteChromeAdapter", "เลือก legacy หรือ console แบบ explicit"),
        ("Pilot coverage", "/ot-employee ผ่าน AppShell; /ot-summary ยัง legacy"),
        ("Migration plan", "ต่อไป /ot-summary -> BEML -> purchasing-all -> purchasing -> home -> Shop Order"),
    ]
    add_table(doc, ["ส่วน", "พฤติกรรมปัจจุบัน"], shell_rows,
              [2000, 7360], header_fill=COLORS["light_green"], font_size=9.3, first_col_bold=True)
    add_callout(doc, "ข้อสังเกต", "เอกสาร redesign ระบุเป้าหมายให้ครบ 7 routes แต่ code snapshot ยังเปิด shell เพียง pilot route. ผู้ดูแลควรถือว่างาน migrate UI ยังไม่ complete แม้ production build จะผ่าน.", fill=COLORS["light_rose"], accent=COLORS["rose"])

    add_heading(doc, "Responsive และ accessibility", 2)
    add_bullet(doc, "Desktop ใช้ Sidebar; mobile ใช้ drawer พร้อม overlay, Escape, focus เข้า close button, focus กลับ trigger และ body-scroll lock.")
    add_bullet(doc, "ทุก route ใช้ Next Link สำหรับ internal navigation และ current route สื่อด้วย aria-current พร้อม cue มากกว่าสีอย่างเดียว.")
    add_bullet(doc, "globals.css กำหนด Prompt, EGAT blue/amber/green/rose, focus ring 3px และ prefers-reduced-motion fallback.")


def add_page_walkthrough(doc: Document):
    add_kicker(doc, "04  /  PAGE WALKTHROUGH", COLORS["blue"])
    add_heading(doc, "การทำงานของแต่ละหน้า", 1)

    add_heading(doc, "4.1 หน้าหลัก /: Work Order Operations", 2)
    add_body(doc, "หน้า Home เรียก /api/dashboard และรีเฟรชอัตโนมัติทุก 30 วินาที พร้อมตัวเลือกปี/เดือนที่อัปเดตค่า filter ในชีตเมื่อผู้ใช้เลือกใหม่. หน้าแสดง status W/O (SAP, Pending, Finish), กลุ่ม W11-W14, ยอดรวม, equipment breakdown และ Load Factor/Man.")
    add_table(doc, ["ชุดข้อมูล", "การนำเสนอ"], [
        ("Operation status", "pie chart + total + progress cards สำหรับ Pending, Finish, SAP"),
        ("W/O by group", "การ์ด W11-W14 พร้อม entrance/left/finish/other finish/out"),
        ("Work by group", "Highcharts column + ตาราง equipment ต่อ W11-W14"),
        ("Load Factor / Man", "4 gauge ต่อกลุ่ม: พนักงาน/ลูกจ้าง ปกติ และ +OT"),
    ], [2400, 6960], header_fill=COLORS["light_blue"], font_size=9.3, first_col_bold=True)

    add_heading(doc, "4.2 Purchasing และ Purchasing All", 2)
    add_body(doc, "app/purchasing/page.tsx เป็น reusable PurchasingPageContent ที่รับ apiPath, title, fixedFilters, gauge panel, table column count และ color theme. /purchasing ใช้ข้อมูลสรุปแบบ gold theme; /purchasing-all reuse component เดิมด้วย teal theme และ fixed filters.")
    add_bullet(doc, "กราฟหมวดการซื้อ/จ้างเป็น donut 3D; กราฟสถานะเป็น column 3D ที่คลิกเพื่อ filter ตาราง.")
    add_bullet(doc, "ค้นหาครอบคลุมค่าใน purchase row ทั้งหมด; filter status ใช้ normalization เพื่อลดปัญหาช่องว่าง/คำนำหน้าตัวเลข.")
    add_bullet(doc, "หน้าโหลดข้อมูลซ้ำทุก 30 วินาที และเรียก /api/purchasing หรือ /api/purchasing-all แบบ no-store.")
    add_callout(doc, "ข้อควรระวังด้าน concurrency", "ปี/เดือนของ purchasing ถูกเขียนลงเซลล์ filter กลางใน Google Sheet และรอสูตรคำนวณประมาณ 4 วินาที; ผู้ใช้หลายคนอาจกระทบ filter ของกันและกัน.", fill=COLORS["light_amber"], accent=COLORS["orange"])

    add_heading(doc, "4.3 BEML Inventory", 2)
    add_body(doc, "หน้า /beml-inventory เรียก /api/beml-inventory ซึ่งอ่าน tab PrintCheck และ map หัวคอลัมน์ตามชื่อภาษาไทย/อังกฤษ. หากไม่มี GOOGLE_BEML_INVENTORY_SHEET_ID จะใช้ deterministic demo data 76 รายการและแสดง warning ชัดเจน.")
    add_table(doc, ["มุมมอง", "ความสามารถ"], [
        ("Dashboard", "Bin map แบบคลิกดูรายละเอียด, KPI ทั้งหมด/ปกติ/ใกล้หมด/หมด และ health ring"),
        ("Inventory", "ตาราง desktop และ mobile cards พร้อม sort ตาม code, P/N, name, system, balance, MIN"),
        ("Low stock / Out stock", "tab และ status filter สำหรับรายการต่ำกว่า MIN หรือ balance เป็น 0"),
        ("Detail modal", "แสดง balance, MIN, MAX, gap ต่ำกว่า MIN, action และคำแนะนำ"),
    ], [2400, 6960], header_fill=COLORS["light_green"], font_size=9.3, first_col_bold=True)

    add_heading(doc, "4.4 OT Summary และ OT Employee", 2)
    add_body(doc, "ทั้งสองหน้าใช้ OtSummaryContent จาก app/ot-summary/page.tsx และเรียก /api/ot-summary. query workerType ทำให้ API โหลดข้อมูลเฉพาะ employee หรือ contractor ได้ แม้หน้า legacy default จะเป็น contractor และหน้า /ot-employee จะเป็น employee.")
    add_bullet(doc, "แยกตารางตาม W11-W14 และมี all-total section.")
    add_bullet(doc, "ข้อมูล contractor มี holidayHours, 1x, 1.5x, total money และ 3x; employee มี total และ total2.")
    add_bullet(doc, "ETAS scan tables และ Check OT Error tables แสดงต่อกลุ่ม โดยใช้ name-to-group map ช่วยจัดกลุ่ม error ให้ตรงกับ summary.")
    add_bullet(doc, "source-sheet links ช่วยเปิดชีตต้นทางจากหน้าเว็บโดยตรง.")

    add_heading(doc, "4.5 Shop Order", 2)
    add_body(doc, "Shop Order เป็นหน้าที่มี interaction สูงสุด: ตารางและ pagination อยู่ซ้าย, summary rail อยู่ขวาบน desktop; บน mobile summary จะขึ้นก่อนตาราง. ผู้ใช้ค้นหา/กรอง, เปิดรายละเอียด, เพิ่ม, แก้ไข, ลบ และแนบไฟล์ได้.")
    add_table(doc, ["ส่วน UI", "พฤติกรรม"], [
        ("Toolbar", "query, ปี พ.ศ., เดือน, สถานะ, refresh, clear และเพิ่ม"),
        ("Table", "sort ใหม่สุดก่อน, sticky header, horizontal scroll, 20 รายการต่อหน้า"),
        ("Summary", "ทั้งหมด/รอดำเนินการ/เสร็จสิ้น, doughnut chart และ top 6 receiving units"),
        ("Form", "CustomSelect, CustomDatePicker, drag/drop file, image preview และ progress"),
        ("Detail", "thumbnail ผ่าน authenticated proxy, edit, delete confirmation"),
    ], [2000, 7360], header_fill=COLORS["light_amber"], font_size=9.2, first_col_bold=True)


def add_architecture(doc: Document):
    add_section_break(doc)
    add_kicker(doc, "05  /  ARCHITECTURE & DATA", COLORS["blue"])
    add_heading(doc, "สถาปัตยกรรมและแหล่งข้อมูล", 1)
    add_body(doc, "ระบบใช้โครงสร้าง route-centric: page client state ทำหน้าที่แสดงผลและ filter; API Route Handler ทำหน้าที่อ่าน/เขียน/normalize; lib/googleSheet.ts และ lib/shop-order/repository.ts ทำหน้าที่เป็น server-side integration boundary.")

    add_heading(doc, "แผนภาพการไหลของข้อมูล", 2)
    diagram = doc.add_table(rows=1, cols=3)
    set_table_geometry(diagram, [3000, 3000, 3360], indent_dxa=120, borders="none")
    diagram_text = [
        ("Browser UI", "7 pages\nfilters, charts, tables\ndialogs, upload progress", COLORS["light_blue"], COLORS["blue"]),
        ("Next.js server", "Route Handlers\ndata mapping\ndomain + repository", COLORS["light_amber"], COLORS["orange"]),
        ("Google services", "Sheets Service Account\nDrive OAuth drive.file\nVercel Cron cleanup", COLORS["light_green"], COLORS["green"]),
    ]
    for cell, (title, text, fill, accent) in zip(diagram.rows[0].cells, diagram_text):
        set_cell_shading(cell, fill)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(title)
        set_run_font(r, size=12, color=accent, bold=True)
        for line in text.split("\n"):
            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.space_after = Pt(1)
            rr = p2.add_run(line)
            set_run_font(rr, size=9.5, color=COLORS["ink"])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    add_body(doc, "ข้อมูลทั่วไปใช้ Google Sheets เป็น source of truth. ไฟล์แนบ Shop Order มีเส้นทางเฉพาะ: server เปิด session, browser ส่ง bytes ผ่าน upload proxy ไป Google Drive, แล้ว server verify/finalize ก่อนบันทึก canonical URL ลงชีต.", align=WD_ALIGN_PARAGRAPH.LEFT)

    add_heading(doc, "API และ data source matrix", 2)
    api_rows = [
        ("GET /api/dashboard", "Dashboard W10 All + Dashboard W10 All info", "status, groupStats, equipment, gauges"),
        ("GET /api/purchasing", "Dashboard W10 All + info", "gauges, charts, purchaseList"),
        ("GET /api/purchasing-all", "Dashboard W11 PRPO infoAll", "charts, purchaseList แบบ all"),
        ("GET /api/beml-inventory", "PrintCheck ของ BEML sheet", "inventory rows + systems; demo fallback"),
        ("GET /api/ot-summary", "OT employee/contractor + ETAS + Check OT Error", "employees, contractors, ETAS, errors, totals"),
        ("GET /api/shop-order", "Order1 + DepartmentList + ReceiverList", "orders, departments, receivers"),
        ("POST /api/shop-order/upload-session", "Google Drive OAuth", "fileId, uploadUrl, expiresAt"),
        ("GET /api/shop-order/cleanup", "Google Drive OAuth + Vercel Cron", "aggregate cleanup counters"),
    ]
    add_table(doc, ["Endpoint", "แหล่งข้อมูล", "ผลลัพธ์หลัก"], api_rows,
              [2500, 3600, 3260], header_fill=COLORS["light_blue"], font_size=8.6)

    add_heading(doc, "Google Sheets adapter", 2)
    add_bullet(doc, "getSheetsClient ใช้ JWT Service Account และ scope spreadsheets เท่านั้น.")
    add_bullet(doc, "lib/googleSheet.ts มี helper แยกสำหรับ dashboard, purchasing, OT, ETAS และ BEML inventory.")
    add_bullet(doc, "ข้อมูลจำนวนมากถูก parse จากตำแหน่ง index ของ range ที่กำหนดตายตัว จึงต้องควบคุม schema ของชีตและห้ามย้ายคอลัมน์โดยไม่แก้ mapping.")
    add_bullet(doc, "date values ของ Shop Order ใช้ Google serial number และถูกแปลงเป็น ISO/Gregorian ก่อนแสดงเป็น พ.ศ.")


def add_shop_order_deep_dive(doc: Document):
    add_kicker(doc, "06  /  SHOP ORDER DEEP DIVE", COLORS["blue"])
    add_heading(doc, "Shop Order เชิงลึก", 1)
    add_body(doc, "Shop Order เป็นโมดูลที่มีการแบ่ง boundary ชัดที่สุดในโครงการ: domain utilities, file rules, OAuth classification, upload client, repository, route handlers และ UI components แยกกัน ทำให้ทดสอบได้ตั้งแต่ pure function จนถึง integration-style repository.")

    add_heading(doc, "6.1 โครงสร้างชีต Order1 A-K", 2)
    schema = [
        ("A", "no", "stable sequence; append แล้วอ่าน row จริงก่อนกำหนดเลข"),
        ("B", "from", "server-enforced เป็น หสบ-ช."),
        ("C", "to", "ต้องพบใน DepartmentList"),
        ("D", "number", "ตัวเลข ASCII 6 หลัก"),
        ("E", "dateIn", "วันที่รับ; เขียนเป็น serial + format dd/MM/yyyy"),
        ("F", "subject", "เรื่อง; ต้องไม่ว่าง"),
        ("G", "receivingUnit", "หน่วยงานรับ; ใช้คำนวณ popular units"),
        ("H", "receiverName", "ผู้รับ; มี ReceiverList เป็น suggestion"),
        ("I", "dateOut", "วันที่ออก; มีค่า = เสร็จสิ้น, ว่าง = รอดำเนินการ"),
        ("J", "note", "หมายเหตุ"),
        ("K", "fileUrl", "canonical Drive webViewLink หรือว่าง"),
    ]
    add_table(doc, ["คอลัมน์", "ฟิลด์", "กติกา"], schema,
              [1000, 1800, 6560], header_fill=COLORS["light_blue"], font_size=9.1, first_col_bold=True)

    add_heading(doc, "6.2 CRUD flow", 2)
    add_number(doc, "ผู้ใช้กรอกข้อมูลและเลือกไฟล์; client ตรวจรูปแบบไฟล์และแสดง preview ก่อนเรียก mutation.")
    add_number(doc, "POST upload-session ส่งเฉพาะ orderNumber, name, mimeType, size; server สร้าง Drive file ID และ resumable session.")
    add_number(doc, "Browser PUT ไฟล์ไป /api/shop-order/upload-proxy; upload client รายงาน progress และ retry เฉพาะ network/timeout/429/5xx สูงสุด 3 attempts.")
    add_number(doc, "POST/PATCH /api/shop-order ส่ง order และ uploadedFileId แบบ optional; repository ตรวจ department, date, file metadata, parent, lifecycle และ leading bytes.")
    add_number(doc, "เมื่อผ่านจึงเปิด permission anyone/reader, เปลี่ยน lifecycle เป็น active และเขียน URL ลง Order1; ถ้าไฟล์ไม่ผ่านจะบันทึก order ต่อโดย K ว่างและคืน partial-success.")
    add_number(doc, "DELETE จะ clear A-K ของ row เดิม ไม่ลบ row เพื่อหลีกเลี่ยง sequence shift จาก concurrent request; ไฟล์ OAuth เดิมถูก schedule-delete 30 วัน.")

    add_heading(doc, "6.3 Attachment lifecycle", 2)
    lifecycle = [
        ("pending", "สร้าง session แล้ว", "pendingSince", "cleanup เมื่ออายุ >= 24 ชั่วโมง"),
        ("active", "verify + permission + Sheet link สำเร็จ", "finalizedAt", "ใช้งานเป็นไฟล์แนบปัจจุบัน"),
        ("scheduled_delete", "ไฟล์ถูกแทนที่หรือลบออเดอร์", "deleteAfter + reason", "cleanup ย้ายเข้า Trash เมื่อครบ 30 วัน"),
    ]
    add_table(doc, ["สถานะ", "เกิดเมื่อ", "metadata", "การจัดการ"], lifecycle,
              [1800, 3200, 2400, 1960], header_fill=COLORS["light_amber"], font_size=8.8, first_col_bold=True)
    add_callout(doc, "Legacy boundary", "ระบบไม่ล้างและไม่แก้ไฟล์เดิมจาก Apps Script/โฟลเดอร์ Picture. เฉพาะไฟล์ที่ OAuth app สร้างและมี parent/lifecycle ตรงกับโฟลเดอร์ที่กำหนดจึงจะถูกจัดการ.", fill=COLORS["light_green"], accent=COLORS["green"])

    add_heading(doc, "6.4 File rules", 2)
    add_bullet(doc, "อนุญาต JPEG/JPG, PNG, WebP และ PDF เท่านั้น.")
    add_bullet(doc, "ขนาดต้องมากกว่า 0 และไม่เกิน 10 MB; 10 MB พอดีถือว่าผ่าน.")
    add_bullet(doc, "extension, MIME type และ magic bytes ต้องสอดคล้องกัน; server ตรวจซ้ำผ่าน leading-byte range.")
    add_bullet(doc, "ชื่อไฟล์ user ถูก sanitize และชื่อที่เก็บจริงใช้ SO-{order}-{timestamp}-{shortId}.{extension}; ไม่เก็บ original filename.")
    add_bullet(doc, "thumbnail proxy จำกัด MIME เป็น image/jpeg/png/webp, ใช้ no-store และจำกัด response ไม่เกิน 2 MiB.")

    add_heading(doc, "6.5 API contract", 2)
    contract = [
        ("GET", "/api/shop-order", "load bootstrap", "200 { ok:true, data }"),
        ("POST", "/api/shop-order", "create order + optional fileId", "201 mutation result"),
        ("PATCH", "/api/shop-order", "update by no + optional fileId", "200 mutation result"),
        ("DELETE", "/api/shop-order", "clear order by no", "200 { no }"),
        ("POST", "/api/shop-order/upload-session", "validate metadata + open session", "201 session"),
        ("PUT", "/api/shop-order/upload-proxy", "proxy bytes to Google URL", "200 or safe upload error"),
        ("GET", "/api/shop-order/attachment-thumbnail?no=N", "verified thumbnail", "200 image or 404"),
        ("GET", "/api/shop-order/cleanup", "cron cleanup", "Bearer secret + aggregate counters"),
    ]
    add_table(doc, ["Method", "Path", "หน้าที่", "ผลลัพธ์"], contract,
              [800, 3200, 3100, 2260], header_fill=COLORS["light_blue"], font_size=8.6)


def add_security(doc: Document):
    add_section_break(doc)
    add_kicker(doc, "07  /  SECURITY & RESILIENCE", COLORS["blue"])
    add_heading(doc, "ความปลอดภัยและความทนทาน", 1)
    add_body(doc, "แนวทางของ Shop Order แยกความลับระหว่าง Sheets กับ Drive: Service Account มี scope spreadsheets; Drive ใช้ OAuth ของบัญชีเจ้าของเพียงบัญชีเดียวและ scope drive.file. ฝั่ง client ไม่ได้รับ client secret, refresh token หรือ private key.")
    security_rows = [
        ("Input validation", "JSON content type, same-origin Origin check, 6-digit order, allowed department, strict date, MIME/extension/size/signature"),
        ("Secret handling", "env-only; setup utility ไม่เขียน .env และไม่พิมพ์ access token; error response ไม่ส่ง Google body"),
        ("URL allowlist", "upload session exact https://www.googleapis.com; thumbnail เฉพาะ *.googleusercontent.com; canonical Drive link เท่านั้น"),
        ("Attachment ownership", "ตรวจ file id, parent folder, appProperties, order number, metadata, trashed=false และ leading bytes"),
        ("Logging", "Shop Order internal error log มี operation/category/correlationId; response ใช้ข้อความทั่วไป"),
        ("Cleanup auth", "Bearer secret ตรวจด้วย timingSafeEqual และคืน aggregate counters เท่านั้น"),
        ("Traffic control", "ระบบยัง public-by-link จึงต้องใช้ WAF rate limit mutation 30 requests/IP/10 นาที"),
    ]
    add_table(doc, ["พื้นที่ควบคุม", "มาตรการที่พบใน code"], security_rows,
              [2100, 7260], header_fill=COLORS["light_green"], font_size=9.0, first_col_bold=True)

    add_heading(doc, "Public-access boundary", 2)
    add_callout(doc, "ความเสี่ยงเชิงผลิตภัณฑ์", "ไม่มีหน้า login และผู้ที่เข้าถึง Production URL สามารถเรียก mutation API ได้. Same-origin check ช่วยลด browser CSRF บางกรณี แต่ไม่ใช่ authorization. URL governance และ WAF จึงเป็น defense-in-depth ที่จำเป็น.", fill=COLORS["light_rose"], accent=COLORS["rose"])
    add_body(doc, "สิ่งที่ควรทำใน Production คือจำกัดการเผยแพร่ URL, ตั้ง WAF ตาม path/method ที่กำหนด, ไม่ใส่ secret ใน URL/log, ติดตาม Vercel runtime logs และพิจารณา authentication/authorization หากข้อมูลหรือความเสี่ยงทางธุรกิจเพิ่มขึ้น.")

    add_heading(doc, "Failure modes ที่รองรับ", 2)
    add_bullet(doc, "Google Drive reauth/quota/folder/access/unavailable ถูก map เป็น safe error code ภาษาไทย.")
    add_bullet(doc, "อัปโหลดไฟล์ล้มเหลวแต่บันทึก order สำเร็จได้ และ UI เปิดทางให้เพิ่มไฟล์อีกครั้ง.")
    add_bullet(doc, "Sheet append/update/clear ล้มเหลวหลัง finalize จะพยายาม restore pending metadata และลบ permission ใหม่แบบ best effort.")
    add_bullet(doc, "Cleanup เป็น idempotent: ไม่พบไฟล์หรือไฟล์ถูกย้าย Trash แล้วไม่นับเป็น failure ที่ทำให้ทั้งงานหยุด.")
    add_bullet(doc, "thumbnail ใช้ authenticated proxy และ fallback เมื่อไม่พบรูปย่อ/ชนิดข้อมูลไม่ปลอดภัย/ขนาดเกิน.")


def add_quality(doc: Document):
    add_section_break(doc)
    add_kicker(doc, "08  /  QUALITY", COLORS["blue"])
    add_heading(doc, "การทดสอบและผลตรวจ", 1)
    add_body(doc, "โครงการมี test coverage หนาแน่นที่สุดใน Shop Order และ navigation shell. ผลรันจริงของ snapshot นี้ช่วยยืนยันว่า behavior สำคัญผ่าน แต่ยังไม่เท่ากับการรับรอง Production แบบ end-to-end เพราะ browser manual acceptance และ live Google integration ไม่ได้รันในรอบนี้.")
    test_rows = [
        ("Domain", "วันที่, status, search/filter, sort, pagination, summary, 10,000-row stress shape"),
        ("File rules", "MIME/extension/size/signature, sanitize filename, truncated bytes"),
        ("Drive OAuth", "env, scope drive.file, error classification"),
        ("Repository", "A-K, append row, stable sequence, lifecycle, cleanup, thumbnail, compensation"),
        ("API handlers", "JSON, origin, validation, safe envelope, cron auth, proxy"),
        ("UI interaction", "navigation, shell drawer, PageHeader, dialogs, filters, upload failure"),
        ("Setup script", "loopback callback, state, offline consent, refresh token, folder creation"),
    ]
    add_table(doc, ["กลุ่ม test", "สิ่งที่ตรวจ"], test_rows,
              [2200, 7160], header_fill=COLORS["light_blue"], font_size=9.2, first_col_bold=True)

    add_heading(doc, "ผลตรวจคำสั่ง", 2)
    result_rows = [
        ("PASS", "npm run test:unit", "Vitest 4.1.10: 25 test files, 270 tests"),
        ("PASS", "npm run test:oauth-setup", "Node test runner: 5 tests"),
        ("PASS", "npx tsc --noEmit", "TypeScript strict check"),
        ("PASS", "npm run build", "Next.js 16.2.12, Turbopack, 10 static pages + dynamic APIs"),
        ("PASS", "npx eslint components/layout components/navigation app/layout.tsx app/ot-summary/page.tsx app/ot-employee/page.tsx", "ส่วน shell/navigation/pilot ที่ตรวจแยก"),
        ("FAIL", "npm run lint", "154 problems: 144 errors, 10 warnings"),
    ]
    add_table(doc, ["สถานะ", "คำสั่ง", "ผลตรวจ"], result_rows,
              [1000, 4800, 3560], header_fill=COLORS["light_amber"], font_size=8.8)

    add_heading(doc, "รายละเอียด lint baseline", 2)
    add_body(doc, "Full lint fail จากหลายกลุ่ม: explicit any ใน API/Google adapter, setState-in-effect ในบางหน้าและ Shop Order UI, unescaped quotes ใน BEML, unused vars ใน purchasing-all และ require() ใน scratch utilities. ปัญหาเหล่านี้ถูกบันทึกเป็น technical debt; ไม่ได้แก้ในรอบการจัดทำรูปเล่ม เพราะคำขอนี้เป็นงานอ่าน/จัดทำเอกสาร ไม่ใช่งานแก้ source.")
    add_heading(doc, "Browser verification handoff", 2)
    add_body(doc, "เอกสาร verification วันที่ 31 กรกฎาคม 2026 ระบุว่า production server รันได้, /ot-employee และ /ot-summary ตอบ HTTP 200 และ automated checkpoint ผ่านที่ commit 3f3cbc6 แต่ browser automation ใน environment ถูก block ด้วย sandbox ACL และไม่มี Playwright/Puppeteer/Cypress. ดังนั้น viewport 360x800 ถึง 1366x768, focus/drawer, Back/Forward และ console diagnostics ยังเป็นรายการ pending.")
    add_callout(doc, "เกณฑ์การรายงาน", "PASS ในส่วนนี้หมายถึงคำสั่งที่รันและผ่านจริง; PENDING หมายถึงยังไม่มีหลักฐานจาก browser/live service; ไม่ควรตีความว่า Production พร้อมสมบูรณ์โดยอัตโนมัติ.", fill=COLORS["light_amber"], accent=COLORS["orange"])


def add_findings(doc: Document):
    add_kicker(doc, "09  /  FINDINGS", COLORS["blue"])
    add_heading(doc, "ข้อค้นพบและงานค้าง", 1)
    add_body(doc, "รายการต่อไปนี้เป็นข้อสังเกตจากการอ่าน code และเอกสาร ไม่ใช่การแก้ไขในรอบนี้. การจัดระดับใช้ผลกระทบต่อความปลอดภัย ความถูกต้องของข้อมูล และความพร้อมดูแลระบบ.")
    finding_rows = [
        ("สูง", "Public mutation API ไม่มี login", "ผู้เข้าถึง URL เรียก CRUD ได้; ต้องคุม URL/WAF หรือวาง auth ก่อนขยายการใช้งาน"),
        ("สูง", "Dashboard API คืน debugInfo 10 แถวแรก", "app/api/dashboard/route.ts ส่ง rawData.slice(0, 10) กลับ client; อาจเผยข้อมูลชีตที่ไม่ควรอยู่ใน response"),
        ("กลาง", "full lint ไม่ผ่าน 144 errors/10 warnings", "เพิ่มความเสี่ยง regression และทำให้ CI gate ที่ใช้ lint ไม่ผ่าน"),
        ("กลาง", "UI shell migrate เพียง /ot-employee", "เป้าหมาย 7 routes ยังไม่ครบ; visual/interaction vocabulary ยังต่างกันระหว่างหน้า"),
        ("กลาง", "filter กลางใน Google Sheet", "dashboard/purchasing เขียนปี/เดือนลงเซลล์กลางและรอสูตร; concurrent users อาจเห็นข้อมูลตามคนอื่น"),
        ("กลาง", "env contract ยังไม่เป็นชุดเดียว", ".env.example ไม่มี BEML/OT vars ที่ code ใช้; .env.local snapshot ไม่มี OAuth/cron vars ที่ Shop Order B1 ต้องใช้"),
        ("ต่ำ", "hard-coded sheet IDs/ranges", "การเปลี่ยน tab/column ต้องแก้ code และเพิ่ม operational runbook/contract test"),
        ("ต่ำ", "scratch ถูก lint รวม", "utilities แบบ CommonJS ที่ไม่ใช่ runtime ทำให้ lint noise สูง; ควรแยก scope หรือปรับ config"),
    ]
    add_table(doc, ["ระดับ", "ข้อค้นพบ", "ผลกระทบ/คำแนะนำ"], finding_rows,
              [900, 3000, 5460], header_fill=COLORS["light_rose"], font_size=8.6)

    add_heading(doc, "ข้อเสนอแนะตามลำดับ", 2)
    add_number(doc, "แก้หรือเอา debugInfo ออกจาก public response และทบทวน error.message ของ legacy API ให้ใช้ safe error envelope แบบเดียวกับ Shop Order.")
    add_number(doc, "กำหนด access model ให้ชัด: ถ้ายัง public ให้บังคับ WAF/URL governance; ถ้าข้อมูลมีความลับ ให้เพิ่ม authentication/authorization.")
    add_number(doc, "จัดทำ env contract กลาง แยกกลุ่ม Dashboard, OT, BEML และ Shop Order พร้อมตรวจชื่อ GOOGLE_SHEET_ID กับ SHOP_ORDER_SHEET_ID.")
    add_number(doc, "ทำ lint baseline cleanup เป็นงานแยก: ย้าย scratch ออกจาก lint, ลด any, แก้ setState-in-effect และ unescaped entities.")
    add_number(doc, "ทำ browser verification matrix ให้ครบทุก route/viewport และทดสอบ live Sheet/Drive ด้วย test order ที่ลบ/กู้ได้.")
    add_number(doc, "เดินหน้า UI shell migration ตามแผน โดยรักษา route data behavior เดิมและเพิ่ม regression test ทุก route.")
    add_callout(doc, "ไม่ควรทำพร้อมกัน", "อย่าเปิด Cache Components/Partial Prefetching หรือเปลี่ยน Next.js ระหว่าง shell migration โดยไม่มี version-matched guide และ data freshness/privacy review.", fill=COLORS["light_amber"], accent=COLORS["orange"])


def add_runbook(doc: Document):
    add_kicker(doc, "10  /  RUNBOOK", COLORS["blue"])
    add_heading(doc, "คู่มือ deploy และ runbook", 1)
    add_heading(doc, "10.1 เริ่มต้นในเครื่อง", 2)
    add_code_block(doc, [
        "npm install",
        "Copy-Item .env.example .env.local",
        "npm run dev",
        "# เปิด http://localhost:3000/shop-order",
    ])
    add_body(doc, "ให้ใช้ Node.js/npm ตาม package-lock และตั้ง environment ให้ครบตาม capability ที่จะใช้งาน. อย่า commit .env.local, private key, OAuth client secret หรือ refresh token.")

    add_heading(doc, "10.2 Environment matrix", 2)
    env_rows = [
        ("GOOGLE_CLIENT_EMAIL", "Sheets", "Service Account email"),
        ("GOOGLE_PRIVATE_KEY", "Sheets", "private key แบบมี \\n ได้"),
        ("GOOGLE_SHEET_ID", "Dashboard/legacy", "main dashboard sheet ที่ lib/googleSheet.ts ใช้"),
        ("SHOP_ORDER_SHEET_ID", "Shop Order", "spreadsheet ของ Order1"),
        ("SHOP_ORDER_SHEET_NAME", "Shop Order", "ปกติคือ Order1"),
        ("GOOGLE_OT_EMPLOYEE_SHEET_ID", "OT employee", "source spreadsheet"),
        ("GOOGLE_OT_CONTRACTOR_SHEET_ID", "OT contractor", "source spreadsheet"),
        ("GOOGLE_BEML_INVENTORY_SHEET_ID", "BEML", "ถ้าไม่ตั้งค่าจะเข้า demo mode"),
        ("GOOGLE_DRIVE_OAUTH_CLIENT_ID", "Shop Order Drive", "OAuth client"),
        ("GOOGLE_DRIVE_OAUTH_CLIENT_SECRET", "Shop Order Drive", "server secret"),
        ("GOOGLE_DRIVE_OAUTH_REFRESH_TOKEN", "Shop Order Drive", "refresh token ของเจ้าของ"),
        ("SHOP_ORDER_DRIVE_FOLDER_ID", "Shop Order Drive", "Picture-OAuth folder"),
        ("SHOP_ORDER_CRON_SECRET / CRON_SECRET", "Cleanup", "ต้องตรงกันสำหรับ Vercel Cron"),
    ]
    add_table(doc, ["ตัวแปร", "ความสามารถ", "หน้าที่"], env_rows,
              [3300, 2200, 3860], header_fill=COLORS["light_blue"], font_size=8.6)

    add_heading(doc, "10.3 ตั้งค่า Drive OAuth B1", 2)
    for step in [
        "เปิด Google Drive API และตั้ง OAuth consent ให้บัญชี w10egat.project@gmail.com ใช้งานได้.",
        "สร้าง OAuth Client แบบ Desktop app แล้วตั้ง GOOGLE_DRIVE_OAUTH_CLIENT_ID/SECRET ชั่วคราวใน terminal.",
        "รัน npm run shop-order:setup-drive; เปิด URL ที่ utility แสดงเอง, อนุมัติ drive.file และรอ callback ที่ 127.0.0.1.",
        "คัดลอกเฉพาะ refresh token กับ Picture-OAuth folder ID ไปยัง secret manager/Vercel; ไม่ส่งผ่าน chat, screenshot หรือ log.",
        "ตั้ง SHOP_ORDER_CRON_SECRET และ CRON_SECRET ให้ตรงกัน แล้ว redeploy.",
    ]:
        add_number(doc, step)

    add_heading(doc, "10.4 Deploy verification", 2)
    add_code_block(doc, [
        "npm run test:unit",
        "npm run test:oauth-setup",
        "npx tsc --noEmit",
        "npm run build",
        "npm run lint   # ต้องแก้ baseline ก่อนใช้เป็น hard gate",
    ])
    checklist = [
        ("Application", "เปิดทุก route และตรวจ loading/error/empty state"),
        ("Dashboard", "เปลี่ยนปี/เดือนและตรวจไม่ให้ filter ของผู้ใช้อื่นถูกทับ"),
        ("Shop Order", "สร้างเลข 6 หลักพร้อม PNG/JPEG ขนาดเล็ก, ตรวจ Order1 และ Drive"),
        ("Replacement", "แก้รายการด้วย PDF แล้วตรวจไฟล์เดิมเป็น scheduled_delete"),
        ("Failure", "จำลอง upload fail; order ต้องถูกบันทึกและมีปุ่มเพิ่มไฟล์อีกครั้ง"),
        ("Cleanup", "เรียก cleanup ด้วย Bearer ที่ถูกต้อง; secret ผิดต้องได้ 401"),
        ("Security", "ตรวจ WAF 30 requests/IP/10 นาที และห้าม session URL อยู่ใน logs"),
    ]
    add_table(doc, ["พื้นที่ตรวจ", "Acceptance check"], checklist,
              [2000, 7360], header_fill=COLORS["light_green"], font_size=9.2, first_col_bold=True)

    add_heading(doc, "10.5 Troubleshooting ที่พบบ่อย", 2)
    trouble = [
        ("GET Shop Order 500", "ตรวจ Sheets service account, private key และสิทธิ์ Editor ใน spreadsheet"),
        ("Drive configuration", "ตรวจ OAuth 3 ค่าและ folder ID ของ Picture-OAuth; ไม่ใช้โฟลเดอร์ Picture legacy"),
        ("Drive 403/404", "ตรวจ quota, Drive API, folder ownership และสิทธิ์ของ OAuth client"),
        ("ไม่มี refresh token", "ถอนสิทธิ์แอปแล้วรัน setup ใหม่ด้วย prompt=consent"),
        ("thumbnail ไม่ขึ้น", "กด refresh; ตรวจว่าไฟล์ active, ไม่อยู่ Trash และเป็นชนิดที่รองรับ"),
        ("upload fail", "client retry network/timeout/429/5xx; 4xx จะไม่ retry และ order ยังบันทึกได้"),
    ]
    add_table(doc, ["อาการ", "แนวทาง"], trouble,
              [2400, 6960], header_fill=COLORS["light_amber"], font_size=9.1, first_col_bold=True)


def add_appendix(doc: Document):
    add_section_break(doc)
    add_kicker(doc, "11  /  APPENDIX", COLORS["blue"])
    add_heading(doc, "ภาคผนวก", 1)

    add_heading(doc, "11.1 File map ที่ควรรู้", 2)
    file_rows = [
        ("app/layout.tsx", "Root metadata, Prompt font, ShellMigrationGate"),
        ("app/page.tsx", "Home dashboard client page"),
        ("app/purchasing/page.tsx", "PurchasingPageContent + purchasing UI"),
        ("app/ot-summary/page.tsx", "Shared OT content for contractor/employee"),
        ("components/layout/*", "AppShell, sidebar, drawer, PageHeader, route chrome"),
        ("components/navigation/*", "Canonical destinations + NavigationMenu"),
        ("components/shop-order/*", "Shop Order dashboard, table, form, dialogs"),
        ("lib/googleSheet.ts", "Sheets client and source-specific readers"),
        ("lib/shop-order/repository.ts", "CRUD, Drive session, finalize, thumbnail, cleanup"),
        ("lib/shop-order/domain.ts", "date conversion, filters, sort, summary, pagination"),
        ("lib/shop-order/file-rules.ts", "approved file types and signatures"),
        ("lib/shop-order/attachment-lifecycle.ts", "pending/active/scheduled-delete state"),
        ("scripts/setup-shop-order-drive-oauth.mjs", "local loopback OAuth setup utility"),
        ("vercel.json", "daily cleanup cron at 17 18 * * * UTC"),
    ]
    add_table(doc, ["ไฟล์/กลุ่มไฟล์", "บทบาท"], file_rows,
              [3500, 5860], header_fill=COLORS["light_blue"], font_size=8.8)

    add_heading(doc, "11.2 Glossary", 2)
    glossary = [
        ("Order1", "ชีตหลักของ Shop Order; โครงสร้างข้อมูล A-K"),
        ("ETAS Scan Data", "ข้อมูลสแกนเวลา ใช้ประกอบตาราง OT ของ worker type เดียวกัน"),
        ("drive.file", "Google Drive scope ที่แอปจัดการเฉพาะไฟล์/โฟลเดอร์ที่สร้างหรือได้รับสิทธิ์ผ่านแอป"),
        ("Pending upload", "ไฟล์ที่สร้าง session แล้วแต่ยังไม่ถูกผูกกับ order"),
        ("Partial success", "Sheet order บันทึกสำเร็จ แต่ attachment ไม่สำเร็จ; UI ต้องให้ลองเพิ่มไฟล์ภายหลัง"),
        ("Legacy file", "ไฟล์/โฟลเดอร์เดิมจาก Apps Script ที่ OAuth B1 ไม่แก้และไม่ล้าง"),
        ("no-store", "response/cache policy ที่ใช้ป้องกันข้อมูล dashboard/thumbnail ค้าง"),
        ("WAF", "Web Application Firewall; ในระบบนี้ใช้จำกัด mutation traffic เป็น defense-in-depth"),
    ]
    add_table(doc, ["คำศัพท์", "ความหมายในโครงการ"], glossary,
              [2200, 7160], header_fill=COLORS["light_green"], font_size=9.0, first_col_bold=True)

    add_heading(doc, "11.3 แหล่งข้อมูลที่ใช้ทำรูปเล่ม", 2)
    sources = [
        "README.md, PRODUCT.md, CONTEXT.md และ AGENTS.md",
        "app/, components/, lib/ และ scripts/ ใน repository snapshot",
        "docs/superpowers/specs/ และ docs/superpowers/plans/",
        "docs/superpowers/verification/2026-07-31-task5-browser-verification-handoff.th.md",
        "ผลคำสั่ง npm run test:unit, npm run test:oauth-setup, npx tsc --noEmit, npm run build, npm run lint และ scoped ESLint",
    ]
    for source in sources:
        add_bullet(doc, source)
    add_callout(doc, "ขอบเขตความรับผิดชอบ", "รูปเล่มนี้เป็นเอกสารวิเคราะห์และ runbook จาก code snapshot; ไม่ได้เปลี่ยน API, schema, environment, WAF หรือ source behavior ของโปรเจกต์.", fill=COLORS["light_blue"], accent=COLORS["blue"])
    add_body(doc, "จบเอกสาร", color=COLORS["gray"], italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_contents(doc)
    add_executive_summary(doc)
    add_scope_stack(doc)
    add_routes_shell(doc)
    add_page_walkthrough(doc)
    add_architecture(doc)
    add_shop_order_deep_dive(doc)
    add_security(doc)
    add_quality(doc)
    add_findings(doc)
    add_runbook(doc)
    add_appendix(doc)

    core = doc.core_properties
    core.title = "W10 Dashboard - รูปเล่มวิเคราะห์ระบบและคู่มือปฏิบัติการ"
    core.subject = "Technical and operations project booklet"
    core.author = ""
    core.last_modified_by = ""
    core.comments = "Generated from repository snapshot dated 2026-07-31."

    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
